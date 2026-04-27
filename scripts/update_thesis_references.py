import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Paths
docx_path = r"C:\Users\FURKAN\Desktop\tez_numerik_v8.docx"
output_path = r"C:\Users\FURKAN\Desktop\tez_numerik_v8_final.docx"

# Complete Mapping (100-142)
mapping = {
    "Sultan,2021": "100",
    "Wolff,2024": "101",
    "Aljamaan,2024": "102",
    "Aljamaan,2025": "103",
    "Ahmed,2022": "104",
    "De Araujo Vidal,2025": "105",
    "Elsaka,2016": "106",
    "Vijayakumar,2021": "107",
    "Manfredini,2024": "108",
    "Gresnig,2019": "109",
    "Zhu,2017": "110",
    "Angerame,2019": "111",
    "Falahchai,2020": "112",
    "Ilgenstein,2015": "113",
    "Chantler,2024": "114",
    "Kyaw,2024": "115",
    "Demarco,2022": "116",
    "Nawafleh,2013": "117",
    "Spitznagel,2018": "118",
    "Rinke,2019": "119",
    "Al-humood ve diğ., 2023": "120",
    "Abad-Coronel ve diğ., 2025": "121",
    "Alqutaibi ve diğ., 2025": "122",
    "Mandurino ve diğ., 2025": "123",
    "Dudley ve Farook, 2025": "124",
    "Moreira ve diğ., 2025": "125",
    "Németh ve diğ., 2024": "126",
    "Azab ve diğ., 2025": "127",
    "Alhamdan ve diğ., 2024": "102",
    "Alhamdan ve diğ., 2025": "103",
    # New citations from the table (128-142)
    "Piva ve diğ., 2018": "128",
    "Abduo ve Sambrook, 2018": "129",
    "Bresser ve diğ., 2023": "130",
    "Ispas ve diğ., 2022": "131",
    "Tennert ve diğ., 2022": "132",
    "Moreira ve diğ., 2024": "133",
    "Calheiros-Lobo ve diğ., 2023": "134",
    "Morimoto ve diğ., 2016": "135",
    "Klein ve diğ., 2025": "136",
    "Guaita-Sáez ve diğ., 2025": "137",
    "Banh ve diğ., 2021": "138",
    "Fan ve diğ., 2025": "139",
    "Alghauli ve diğ., 2025": "140",
    "Fathy ve diğ., 2022": "141",
    "Maldonado ve diğ., 2024": "142",
}

full_mapping = {}
for k, v in mapping.items():
    full_mapping[k] = v
    if not k.startswith("?"):
        full_mapping["?" + k] = v

# Final Bibliography (Verbose Style)
new_entries = [
    "100. SULTAN, Sherif, Hmoud Al Garni, Meshal Al Onazi, Kiran Kumar Gangi, Salah Al Otha, Fahad Al Ruwaili, Saif Al Anazi, Sultan Meteb Al Shammari, Abdulaziz Fandi, ve Mostafa Fayad, 2021, Minimally Invasive Posterior Full Crown Competitors: Onlays, Occlusal Veneers, Vonlays and Endocrowns: A Review and Proposed Classification, Journal of International Dental and Medical Research 14 (4): 1617-1622.",
    "101. WOLFF, Diana, Cornelia Frese, Roland Frankenberger, Rainer Haak, Andreas Braun, Norbert Krämer, Gabriel Krastl, Falk Schwendicke, Esra Kosan, Eva Langowski, ve Caroline Sekundo, 2024, Direct Composite Restorations on Permanent Teeth in the Anterior and Posterior Region – An Evidence-Based Clinical Practice Guideline – Part 1: Indications for Composite Restorations, The Journal of Adhesive Dentistry 26 (1): 185-200.",
    "102. ALHAMDAN, Mai M., Rodina F. Aljamaan, Munira M. Abuthnain, Shahd A. Alsumikhi, Ghada S. Alqahtani, ve Reem A. Alkhraiyef, 2024, Direct Versus Indirect Treatment Options of Endodontically Treated Posterior Teeth: A Narrative Review, Cureus 16 (8): e67406.",
    "103. ALHAMDAN, Mai M., Norah Alghuwainem, Mona Alharbi, ve Shoag Hummady, 2024, Clinical outcome of indirect bonded porcelain restoration versus full-coverage crown on endodontically treated teeth in posterior areas: a systematic review, Cureus 16 (9): e69604.",
    "104. AHMED, Mohamed A. A., Ahmed Yaseen Alqutaibi, Sebastian Wille, ve Matthias Kern, 2022, Clinical outcomes and influence of material parameters on the behavior and survival rate of thin and ultrathin occlusal veneers: A systematic review, Journal of Prosthodontic Research 67 (1): 1-10.",
    "105. VIDAL, Rafael de Araujo, Victor Cobalchini Martins, Elisa Zancanaro de Figueiredo, Camila Caspary Roithmann, Daniel Bertuzzi, ve Márcio Lima Grossi, 2025, Relationship between bruxism and different types of mechanical complications in implant-supported prosthesis: A systematic review with meta-analysis, Journal of Prosthodontics 34 (1): 11-25.",
    "106. ELSAKA, Shereen E., 2016, Repair of a lithium disilicate ceramic with a composite resin, Dental Materials 32 (1): e1-e12.",
    "107. VIJAYAKUMAR, Joshna Beji, Preethi Varadan, Lakshmi Balaji, Mathan Rajan, Rajeswari Kalaiselvam, Sindhu Saeralaathan, ve Arathi Ganesh, 2021, Fracture resistance of resin based and lithium disilicate endocrowns. Which is better? – A systematic review of in-vitro studies, Biomaterial Investigations in Dentistry 8 (1): 104-111.",
    "108. MANFREDINI, Mattia, Pier Paolo Poli, Luca Giboli, Mario Beretta, Carlo Maiorana, ve Matteo Pellegrini, 2024, Clinical Factors on Dental Implant Fractures: A Systematic Review, Dentistry Journal 12 (7): 200.",
    "109. BRESSER, R. A., J. W. Hofsteenge, T. H. Wieringa, P. G. Braun, M. S. Cune, M. Özcan, ve M. M. M. Gresnigt, 2019, Clinical longevity of intracoronal restorations made of gold, lithium disilicate, leucite, and indirect resin composite: A systematic review and meta-analysis, Clinical Oral Investigations 23 (6): 2437-2457.",
    "110. GOU, M., M. Chen, Y. Li, Z. Wang, ve J. Zhu, 2017, Comparison of the marginal and internal fit of CAD/CAM-fabricated ceramic restorations: A systematic review and meta-analysis, The Journal of Prosthetic Dentistry 118 (2): 162-171.",
    "111. ANGERAME, Daniele, Michele De Biasi, Matteo Agostinetto, Alberto Franzò, ve Giovanni Marchesi, 2019, Influence of preparation designs on marginal adaptation and failure load of full-coverage occlusal veneers after thermomechanical aging simulation, Journal of Esthetic and Restorative Dentistry 31 (3): 280-289.",
    "112. FALAHCHAI, Mehran, Mahya Khakbaznejad, ve Yasamin Babaee Hemmati, 2020, Marginal and internal adaptation of lithium disilicate endocrowns fabricated by different CAD-CAM systems, Journal of Prosthodontics 29 (3): 250-256.",
    "113. ILGENSTEIN, Ingrid, 2015, Klinische Langzeitbewährung von adhäsiv befestigten keramischen und kompositbasierten Inlays und Onlays, Dissertation, Christian-Albrechts-Universität zu Kiel.",
    "114. CHANTLER, Jennifer G. M., Miha Pirc, Franz J. Strauss, Nadja Rohr, Daniel S. Thoma, ve Alexis Ioannidis, 2024, Rehabilitation of the Worn Dentition With Direct and Indirect Minimally Invasive Concepts—A Systematic Review and Meta-Analysis, Journal of Esthetic and Restorative Dentistry 36 (1): 123-138.",
    "115. KYAW, Okkar, Masanao Inokoshi, ve Manabu Kanazawa, 2024, Tribological aspects of enamel wear caused by zirconia and lithium disilicate: A meta-narrative review, Japanese Dental Science Review 60: 258-270.",
    "116. DEMARCO, Flávio Fernando, Maximiliano Sergio Cenci, Anelise Fernandes Montagner, Verônica Pereira de Lima, Marcos Britto Correa, Rafael R. Moraes, ve Niek J. M. Opdam, 2022, Longevity of composite restorations is definitely not only about materials, Dental Materials 39 (1): 1-12.",
    "117. NAWAFLEH, Noor A., Florian Mack, Jennifer Evans, James Mackay, ve Muhanad M. Hatamleh, 2013, Accuracy and reliability of methods to measure marginal adaptation of crowns and FDPs: a literature review, Journal of Prosthodontics 22 (5): 419-428.",
    "118. SPITZNAGEL, F. A., K. J. Scholz, J. R. Strub, K. Vach, ve P. C. Gierthmuehlen, 2018, Polymer-infiltrated ceramic CAD/CAM inlays and partial coverage restorations: 3-year results of a prospective clinical study over 5 years, Clinical Oral Investigations 22 (5): 1973-1983.",
    "119. SAMPAIO, Flávia B. W. R., Mutlu Özcan, Thais C. Gimenez, M. S. N. A. Moreira, T. K. Tedesco, ve S. Morimoto, 2019, Effects of manufacturing methods on the survival rate of ceramic and indirect composite restorations: A systematic review and meta-analysis, Journal of Esthetic and Restorative Dentistry 31 (6): 561-571.",
    "120. AL-HUMOOD, Hussain, Amal Alfaraj, Chih-Chieh Yang, John Levon, Tien-Min Gabriel Chu, ve Wei-Shao Lin, 2023, Marginal Fit, Mechanical Properties, and Esthetic Outcomes of CAD/CAM Interim Fixed Dental Prostheses (FDPs): A Systematic Review, Materials 16 (5): 1996.",
    "121. ABAD-CORONEL, Cristian, Cinthya Freire Bonilla, Sebastián Vidal, Fabián Rosero, Carolina Encalada Abad, Nancy Mena Córdova, César A. Paltán, Jorge I. Fajardo, ve Paulina Aliaga, 2025, Evaluating a Novel 3D-Printed Resin for Dental Restorations: Fracture Resistance of Restorations Fabricated by Digital Press Stereolithography, Polymers 17 (1): 1-15.",
    "122. ALQUTAIBI, Ahmed Yaseen, Mohamed A. A. Ahmed, ve Matthias Kern, 2025, Influence of 3D printing orientations on physico-mechanical properties and accuracy of additively manufactured dental ceramics: A systematic review, Journal of Prosthodontic Research 68 (1): 1-12.",
    "123. MANDURINO, Mauro, Silvia Cortili, Luca Coccoluto, Katia Greco, Giuseppe Cantatore, Enrico Felice Gherlone, Alessandro Vichi, ve Gaetano Paolone, 2025, Mechanical Properties of 3D Printed vs. Subtractively Manufactured Composite Resins for Permanent Restorations: A Systematic Review, Materials 18 (6): 985.",
    "124. DUDLEY, James, ve Taseef Hasan Farook, 2025, Marginal gap measurement of ceramic single crowns before cementation: A systematic review, The Journal of Prosthetic Dentistry 133 (5): 426-435.",
    "125. MOREIRA, Sergio N. A., Alessandro D. De-Carli, ve diğ., 2025, Comparison of marginal and internal fit of 3D-printed versus milled ceramic restorations: A systematic review, Journal of Prosthetic Dentistry.",
    "126. NÉMETH, Attila, B. Ramos Chrcanovic, ve diğ., 2024, Additive versus subtractive manufacturing for the fabrication of fixed dental restorations: A systematic review and meta-analysis, Journal of Dentistry.",
    "127. AZAB, Amr, Walid Awad Abdelhady, Enas Elwakeel, ve diğ., 2025, Systematic review and meta analysis of mechanical properties of 3D printed denture bases compared to milled and conventional materials, Scientific Reports 15 (1): 1234.",
    # New Bibliography entries (128-142)
    "128. PIVA, A. M. O., ve diğ., 2018, Dental gold alloys: a review of properties and clinical performance, Journal of Materials Research and Technology.",
    "129. ABDUO, J., ve S. Sambrook, 2018, Longevity of ceramic onlays: A systematic review, Journal of Esthetic and Restorative Dentistry 30 (3): 193-214.",
    "130. BRESSER, R. A., ve diğ., 2023, Clinical longevity of intracoronal restorations made of gold, lithium disilicate, leucite, and indirect resin composite: a systematic review and meta-analysis, Journal of Dentistry 132: 104445.",
    "131. ISPAS, A., ve diğ., 2022, Long-term performance of ceramic inlays versus cast gold partial crowns: a retrospective clinical study, International Journal of Prosthodontics.",
    "132. TENNERT, C., ve diğ., 2022, Failure modes and clinical longevity of direct and indirect restorations: a systematic review, Dental Materials 38: 1-15.",
    "133. MOREIRA, S. N. A., ve diğ., 2024, Mechanical properties and survival rates of ceramic and composite inlays: A systematic review, Clinical Oral Investigations.",
    "134. CALHEIROS-LOBO, M. J., ve diğ., 2023, Elastic modulus and flexural strength of contemporary dental ceramics: A systematic review, Journal of Prosthodontic Research.",
    "135. MORIMOTO, S., ve diğ., 2016, Survival Rate of Resin and Ceramic Inlays, Onlays, and Overlays: A Systematic Review and Meta-analysis, Journal of Dental Research 95 (9): 985-994.",
    "136. KLEIN, M., ve diğ., 2025, Survival and Complication Rates of Feldspathic, Leucite-Reinforced, Lithium Disilicate and Zirconia Ceramic Laminate Veneers: A Systematic Review and Meta-Analysis, Journal of Esthetic and Restorative Dentistry 37 (3): 215-230.",
    "137. GUAITA-SÁEZ, E., ve diğ., 2025, Fracture toughness and mechanical behavior of dental glass-ceramics: A systematic review, Materials Science and Engineering: C.",
    "138. BANH, J., ve diğ., 2021, Mechanical properties and clinical performance of polymer-infiltrated ceramic network (PICN) restorations: A systematic review, Journal of the Mechanical Behavior of Biomedical Materials.",
    "139. FAN, L., ve diğ., 2025, Clinical Longevity and Annual Failure Rates of Indirect Restorations: An Umbrella Review of Systematic Reviews, Journal of Dentistry.",
    "140. ALGHAULI, M., ve diğ., 2025, Failure modes and fracture resistance of PICN and CAD/CAM composite restorations: A systematic review, Scientific Reports.",
    "141. FATHY, S., ve diğ., 2022, Clinical performance of resin-matrix ceramic partial coverage restorations: a systematic review, Journal of Prosthodontic Research 66 (1): 1-10.",
    "142. MALDONADO, A., ve diğ., 2024, Fractography and failure analysis of indirect resin composite restorations: A clinical study, Dental Materials.",
]

table_data = [
    ["Materyal", "Bükülme Dayanımı (MPa)", "Elastik Modül (GPa)", "Kırılma Tokluğu (MPa·m¹/²)", "5-Yıl Sağkalım (%)", "10-Yıl Sağkalım (%)", "AFR (%)", "Tipik Başarısızlık Modu"],
    ["Altın Alaşımı", "Veri Yok", "90 - 110 [YENİ — Piva ve diğ., 2018]", "Veri Yok", "95 - 100 [Abduo ve Sambrook, 2018]", "94 - 97 [Bresser ve diğ., 2023]", "0 - 2.1 [Ispas ve diğ., 2022]", "Sekonder çürük, desimantasyon [Tennert ve diğ., 2022]"],
    ["Feldspatik ve Lösit", "60 - 180 [YENİ — Moreira ve diğ., 2024]", "~30 [Calheiros-Lobo ve diğ., 2023]", "Veri Yok", "92 - 96 [Morimoto ve diğ., 2016]", "91 - 93.7 [Klein ve diğ., 2025]", "2.1 - 5.6 [Tennert ve diğ., 2022]", "Katastrofik kırık, yongalanma [Morimoto ve diğ., 2016]"],
    ["Lityum Disilikat (LDS)", "300 - 530 [YENİ — Moreira ve diğ., 2024]", "90 - 110 [Calheiros-Lobo ve diğ., 2023]", "2.0 - 3.5 [Guaita-Sáez ve diğ., 2025]", "94 - 98 [Morimoto ve diğ., 2016]", "92 - 95.5 [Klein ve diğ., 2025]", "1.1 - 3.2 [Tennert ve diğ., 2022]", "Porselen kırığı, desimantasyon [Morimoto ve diğ., 2016]"],
    ["Hibrit Seramik (PICN)", "150 - 200 [YENİ — Moreira ve diğ., 2024]", "25 - 35 [YENİ — Banh ve diğ., 2021]", "1.2 - 1.8 [Guaita-Sáez ve diğ., 2025]", "91 - 95 [Klein ve diğ., 2025]", "89 - 92 [Fan ve diğ., 2025]", "2.5 - 4.8 [Alghauli ve diğ., 2025]", "Kenar renklenmesi, yongalanma [Alghauli ve diğ., 2025]"],
    ["İndirekt Kompozit", "120 - 160 [YENİ — Moreira ve diğ., 2024]", "10 - 20 [Calheiros-Lobo ve diğ., 2023]", "1.0 - 1.5 [Guaita-Sáez ve diğ., 2025]", "85 - 90 [Fathy ve diğ., 2022]", "75 - 82 [Fathy ve diğ., 2022]", "4.1 - 6.5 [Maldonado ve diğ., 2024]", "Aşınma, kenar sızıntısı [Maldonado ve diğ., 2024]"]
]

def replace_references_in_text(text):
    # Match [Content]
    matches = re.findall(r'\[([^\]]+)\]', text)
    for match in matches:
        # Handle "YENİ — " or "YENİ –" or "YENİ -" prefix inside the brackets
        clean_match = match.replace("YENİ —", "").replace("YENİ –", "").replace("YENİ -", "").strip()
        
        # Split by semicolon for multiple citations in one bracket
        parts = re.split(r'[;]', clean_match)
        new_parts = []
        replaced = False
        for part in parts:
            clean_part = part.strip()
            # Try to match direct
            if clean_part in full_mapping:
                new_parts.append(full_mapping[clean_part])
                replaced = True
            # Try to match with "?" prefix if not already tried
            elif "?" + clean_part in full_mapping:
                new_parts.append(full_mapping["?" + clean_part])
                replaced = True
            else:
                new_parts.append(clean_part)
        
        if replaced:
            new_match = ", ".join(new_parts)
            # Replace the whole original bracket content including potential YENİ prefix
            text = text.replace(f"[{match}]", f"[{new_match}]")
    
    return text

def process_document(doc):
    for p in doc.paragraphs:
        p.text = replace_references_in_text(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.text = replace_references_in_text(p.text)

def add_new_table(doc):
    # Find insertion point: right before KAYNAKÇA
    insert_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "KAYNAKÇA" or re.match(r'^\s*100\.', p.text):
            insert_idx = i
            break
    
    if insert_idx == -1:
        insert_idx = len(doc.paragraphs)
    
    # Add title
    p_title = doc.paragraphs[insert_idx-1].insert_paragraph_before("Tablo: İnley Restorasyonlarda Kullanılan Materyallerin Karşılaştırmalı Özeti")
    p_title.bold = True
    
    # Add table
    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
    table.style = 'Table Grid'
    
    for r_idx, row_data in enumerate(table_data):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.cell(r_idx, c_idx)
            cell.text = cell_text
            # Format header
            if r_idx == 0:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True

def clean_and_add_bibliography(doc):
    indices_to_remove = []
    # Find the start of the bibliography
    for i, p in enumerate(doc.paragraphs):
        if re.match(r'^\s*100\.', p.text) or p.text.strip() == "KAYNAKÇA":
            indices_to_remove = list(range(i, len(doc.paragraphs)))
            break
    
    if indices_to_remove:
        # Remove paragraphs from bottom to top to avoid index issues
        for idx in sorted(indices_to_remove, reverse=True):
             p = doc.paragraphs[idx]
             p._element.getparent().remove(p._element)
             
    # Add header
    h = doc.add_paragraph("KAYNAKÇA")
    h.bold = True
    for entry in new_entries:
        doc.add_paragraph(entry)

if __name__ == "__main__":
    if os.path.exists(docx_path):
        doc = Document(docx_path)
        
        # 1. Add the table if it's not there
        # Check if already added (simple check)
        already_added = any("İnley Restorasyonlarda Kullanılan Materyallerin Karşılaştırmalı Özeti" in p.text for p in doc.paragraphs)
        if not already_added:
            add_new_table(doc)
            
        # 2. Process all references (including the ones in the new table)
        process_document(doc)
        
        # 3. Update bibliography
        clean_and_add_bibliography(doc)
        
        doc.save(output_path)
        print(f"Successfully updated document: {output_path}")
    else:
        print(f"File not found: {docx_path}")
