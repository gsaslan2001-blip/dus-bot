from docx import Document

doc = Document(r"C:\Users\FURKAN\Desktop\tez_numerik_v8_updated.docx")
print(f"Total tables: {len(doc.tables)}")

for i, table in enumerate(doc.tables):
    # Check if this table contains the keyword
    found = False
    for row in table.rows:
        for cell in row.cells:
            if "3D Baskı" in cell.text:
                found = True
                break
        if found: break
    
    if found:
        print(f"\n--- Table {i+1} (Found Keyword) ---")
        for row in table.rows[:5]: # Print first 5 rows
            print(" | ".join([cell.text for cell in row.cells]))
    else:
        # Just print first cell to identify
        try:
            print(f"Table {i+1} first cell: {table.rows[0].cells[0].text[:50]}")
        except:
            print(f"Table {i+1} could not read.")
