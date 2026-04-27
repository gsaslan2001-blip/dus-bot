import sys
from docx import Document

if sys.platform == "win32":
    import codecs
    sys.stdout = codecs.getwriter("utf-8")(sys.stdout.detach())

doc = Document(r"C:\Users\FURKAN\Desktop\tez_numerik_v8_final.docx")

print("--- Verification of Table 3 ---")
for table in doc.tables:
    found = False
    for row in table.rows:
        for cell in row.cells:
            if "3D Baskı" in cell.text:
                found = True
                break
        if found: break
    
    if found:
        for row in table.rows:
            text = " | ".join([cell.text for cell in row.cells])
            if "[" in text and "]" in text:
                print(text)

print("\n--- Verification of Bibliography (End) ---")
paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
for p in paragraphs[-30:]:
    print(p)
