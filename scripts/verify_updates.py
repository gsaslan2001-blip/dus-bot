from docx import Document

doc = Document(r"C:\Users\FURKAN\Desktop\tez_numerik_v8_updated.docx")
if len(doc.tables) >= 3:
    table = doc.tables[2]
    print("\n--- Table 3 Content ---")
    for row in table.rows:
        print(" | ".join([cell.text for cell in row.cells]))
else:
    print(f"Only {len(doc.tables)} tables found.")
